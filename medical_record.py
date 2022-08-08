from docx import Document

def generate_medical_record(data):
    symptom_dict = { 1 : "Cough" , 2 : "Muscle aches",  3 : "Tiredness",  4 : "Sore throat",  5 : "Runny nose",  6 : "Stuffy nose",  8 : "Fever ", 9 : "Nause",  10 : "Vomiting", 11 : "Diarrhea", 13 : "Shortness of breath", 14 : " Difficulty in breathing",  15 : " Loss of taste ", 16 : " Loss of smell ", 17 : " Itchy Nose",  18 : " Itchy eyes",  19 : "Itchy inner ear ", 20 : "Sneezing"}
    document = Document()
    document.add_heading('Medical Record', 3)

    document.add_heading('Personal Informations', level=1)
    document.add_paragraph('Name: ' + data['name'])
    document.add_paragraph('Age: ' + data['age'])
    document.add_paragraph('Profession: ' + data['job'])
    document.add_paragraph('Phone Number: ' + data['phone_number'])

    document.add_heading('Symptoms', level=1)
    document.add_paragraph('How long patient is sick: ' + data['days_sick'])
    cont = 0
    if len(data['fever_degree']) > 0:
        for degree in data['fever_degree']:
            p2 = document.add_paragraph("Fever's degree: ")
            if cont == len(data['fever_degree'])-1:
                p2.add_run(degree)
            else:
                p2.add_run(degree + ',')
            cont+=1
    cont = 0
    p1 = document.add_paragraph('Symptoms described by patient: ')
    for symptom in data['symptom_described']:
        if cont == len(data['symptom_described'])-1:
            p1.add_run(symptom)
        else:
            p1.add_run(symptom + ',')
        cont += 1
    p = document.add_paragraph('Symptoms: ')
    for symptom in data['symptoms']:
        p.add_run(symptom_dict.get(symptom) + ',')
    
    document.add_heading('Prediction', level=1)
    document.add_paragraph('Disease predicted by random forest: ' + data['rf_model'])
    document.add_paragraph('Disease predicted by KNN: ' + data['knn_model'])


    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )

    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    # document.add_page_break()
    filename = data['name'] + ".docx"
    document.save(filename)